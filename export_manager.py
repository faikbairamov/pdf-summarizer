"""
Export functionality for multiple formats
"""
import os
import json
import pandas as pd
from typing import List, Dict, Any, Optional
from datetime import datetime
import logging
from pathlib import Path

# For Word document export
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# For PDF export
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    PDF_EXPORT_AVAILABLE = True
except ImportError:
    PDF_EXPORT_AVAILABLE = False

logger = logging.getLogger(__name__)

class ExportManager:
    """Handle export of results to various formats"""
    
    def __init__(self, output_dir: str = "exports"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
    
    def export_to_json(self, results: List[Dict[str, Any]], filename: Optional[str] = None) -> str:
        """Export results to JSON format"""
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"pdf_analysis_{timestamp}.json"
        
        filepath = os.path.join(self.output_dir, filename)
        
        export_data = {
            "export_info": {
                "exported_at": datetime.now().isoformat(),
                "total_documents": len(results),
                "format": "json"
            },
            "results": results
        }
        
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(export_data, f, indent=2, ensure_ascii=False)
        
        logger.info(f"Exported {len(results)} documents to JSON: {filepath}")
        return filepath
    
    def export_to_csv(self, results: List[Dict[str, Any]], filename: Optional[str] = None) -> str:
        """Export results to CSV format"""
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"pdf_analysis_{timestamp}.csv"
        
        filepath = os.path.join(self.output_dir, filename)
        
        # Flatten the results for CSV
        flattened_results = []
        for result in results:
            flat_result = {
                "file_name": result.get("file_name", ""),
                "file_path": result.get("file_path", ""),
                "file_size": result.get("file_size", 0),
                "title": result.get("title", ""),
                "authors": result.get("authors", ""),
                "summary": result.get("summary", ""),
                "keywords": ", ".join(result.get("keywords", [])),
                "character_count": result.get("statistics", {}).get("character_count", 0),
                "word_count": result.get("statistics", {}).get("word_count", 0),
                "sentence_count": result.get("statistics", {}).get("sentence_count", 0),
                "compression_ratio": result.get("statistics", {}).get("compression_ratio", 0),
                "processed_at": result.get("processed_at", "")
            }
            flattened_results.append(flat_result)
        
        df = pd.DataFrame(flattened_results)
        df.to_csv(filepath, index=False)
        
        logger.info(f"Exported {len(results)} documents to CSV: {filepath}")
        return filepath
    
    def export_to_excel(self, results: List[Dict[str, Any]], filename: Optional[str] = None) -> str:
        """Export results to Excel format with multiple sheets"""
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"pdf_analysis_{timestamp}.xlsx"
        
        filepath = os.path.join(self.output_dir, filename)
        
        with pd.ExcelWriter(filepath, engine='openpyxl') as writer:
            # Main results sheet
            flattened_results = []
            for result in results:
                flat_result = {
                    "file_name": result.get("file_name", ""),
                    "file_path": result.get("file_path", ""),
                    "file_size": result.get("file_size", 0),
                    "title": result.get("title", ""),
                    "authors": result.get("authors", ""),
                    "summary": result.get("summary", ""),
                    "keywords": ", ".join(result.get("keywords", [])),
                    "character_count": result.get("statistics", {}).get("character_count", 0),
                    "word_count": result.get("statistics", {}).get("word_count", 0),
                    "sentence_count": result.get("statistics", {}).get("sentence_count", 0),
                    "compression_ratio": result.get("statistics", {}).get("compression_ratio", 0),
                    "processed_at": result.get("processed_at", "")
                }
                flattened_results.append(flat_result)
            
            df_main = pd.DataFrame(flattened_results)
            df_main.to_excel(writer, sheet_name='Analysis Results', index=False)
            
            # Keywords analysis sheet
            all_keywords = []
            for result in results:
                all_keywords.extend(result.get("keywords", []))
            
            if all_keywords:
                from collections import Counter
                keyword_counts = Counter(all_keywords)
                df_keywords = pd.DataFrame([
                    {"keyword": k, "frequency": v} 
                    for k, v in keyword_counts.most_common()
                ])
                df_keywords.to_excel(writer, sheet_name='Keyword Analysis', index=False)
            
            # Statistics summary sheet
            stats_data = {
                "Metric": [
                    "Total Documents",
                    "Average Word Count",
                    "Average Character Count",
                    "Average Compression Ratio",
                    "Total File Size (MB)"
                ],
                "Value": [
                    len(results),
                    df_main['word_count'].mean(),
                    df_main['character_count'].mean(),
                    df_main['compression_ratio'].mean(),
                    df_main['file_size'].sum() / (1024 * 1024)
                ]
            }
            df_stats = pd.DataFrame(stats_data)
            df_stats.to_excel(writer, sheet_name='Statistics Summary', index=False)
        
        logger.info(f"Exported {len(results)} documents to Excel: {filepath}")
        return filepath
    
    def export_to_word(self, results: List[Dict[str, Any]], filename: Optional[str] = None) -> str:
        """Export results to Word document format"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx package is required for Word export")
        
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"pdf_analysis_{timestamp}.docx"
        
        filepath = os.path.join(self.output_dir, filename)
        
        doc = Document()
        
        # Title
        title = doc.add_heading('PDF Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Summary
        doc.add_heading('Summary', level=1)
        summary_para = doc.add_paragraph(f"Total Documents Analyzed: {len(results)}")
        summary_para.add_run(f"\nGenerated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # Individual document results
        for i, result in enumerate(results, 1):
            doc.add_heading(f'Document {i}: {result.get("file_name", "Unknown")}', level=2)
            
            # Document info
            doc.add_heading('Document Information', level=3)
            info_table = doc.add_table(rows=4, cols=2)
            info_table.style = 'Table Grid'
            
            info_table.cell(0, 0).text = 'File Name'
            info_table.cell(0, 1).text = result.get("file_name", "")
            info_table.cell(1, 0).text = 'File Size'
            info_table.cell(1, 1).text = f"{result.get('file_size', 0) / (1024*1024):.2f} MB"
            info_table.cell(2, 0).text = 'Word Count'
            info_table.cell(2, 1).text = str(result.get('statistics', {}).get('word_count', 0))
            info_table.cell(3, 0).text = 'Character Count'
            info_table.cell(3, 1).text = str(result.get('statistics', {}).get('character_count', 0))
            
            # Title
            doc.add_heading('Title', level=3)
            doc.add_paragraph(result.get("title", "No title found"))
            
            # Authors
            doc.add_heading('Authors', level=3)
            doc.add_paragraph(result.get("authors", "No authors found"))
            
            # Summary
            doc.add_heading('Summary', level=3)
            doc.add_paragraph(result.get("summary", "No summary available"))
            
            # Keywords
            doc.add_heading('Keywords', level=3)
            keywords = result.get("keywords", [])
            if keywords:
                doc.add_paragraph(", ".join(keywords))
            else:
                doc.add_paragraph("No keywords found")
            
            # Add page break between documents (except for the last one)
            if i < len(results):
                doc.add_page_break()
        
        doc.save(filepath)
        logger.info(f"Exported {len(results)} documents to Word: {filepath}")
        return filepath
    
    def export_to_pdf(self, results: List[Dict[str, Any]], filename: Optional[str] = None) -> str:
        """Export results to PDF format"""
        if not PDF_EXPORT_AVAILABLE:
            raise ImportError("reportlab package is required for PDF export")
        
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"pdf_analysis_{timestamp}.pdf"
        
        filepath = os.path.join(self.output_dir, filename)
        
        doc = SimpleDocTemplate(filepath, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=1  # Center alignment
        )
        story.append(Paragraph("PDF Analysis Report", title_style))
        story.append(Spacer(1, 12))
        
        # Summary
        story.append(Paragraph("Summary", styles['Heading2']))
        summary_text = f"Total Documents Analyzed: {len(results)}<br/>Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        story.append(Paragraph(summary_text, styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Individual document results
        for i, result in enumerate(results, 1):
            story.append(Paragraph(f"Document {i}: {result.get('file_name', 'Unknown')}", styles['Heading2']))
            
            # Document information table
            info_data = [
                ['File Name', result.get("file_name", "")],
                ['File Size', f"{result.get('file_size', 0) / (1024*1024):.2f} MB"],
                ['Word Count', str(result.get('statistics', {}).get('word_count', 0))],
                ['Character Count', str(result.get('statistics', {}).get('character_count', 0))]
            ]
            
            info_table = Table(info_data, colWidths=[2*inch, 4*inch])
            info_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 14),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(info_table)
            story.append(Spacer(1, 12))
            
            # Title
            story.append(Paragraph("Title", styles['Heading3']))
            story.append(Paragraph(result.get("title", "No title found"), styles['Normal']))
            story.append(Spacer(1, 6))
            
            # Authors
            story.append(Paragraph("Authors", styles['Heading3']))
            story.append(Paragraph(result.get("authors", "No authors found"), styles['Normal']))
            story.append(Spacer(1, 6))
            
            # Summary
            story.append(Paragraph("Summary", styles['Heading3']))
            story.append(Paragraph(result.get("summary", "No summary available"), styles['Normal']))
            story.append(Spacer(1, 6))
            
            # Keywords
            story.append(Paragraph("Keywords", styles['Heading3']))
            keywords = result.get("keywords", [])
            if keywords:
                story.append(Paragraph(", ".join(keywords), styles['Normal']))
            else:
                story.append(Paragraph("No keywords found", styles['Normal']))
            
            story.append(Spacer(1, 20))
        
        doc.build(story)
        logger.info(f"Exported {len(results)} documents to PDF: {filepath}")
        return filepath
    
    def export_all_formats(self, results: List[Dict[str, Any]], base_filename: Optional[str] = None) -> Dict[str, str]:
        """Export results to all available formats"""
        if not base_filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            base_filename = f"pdf_analysis_{timestamp}"
        
        exported_files = {}
        
        # JSON export
        json_file = self.export_to_json(results, f"{base_filename}.json")
        exported_files['json'] = json_file
        
        # CSV export
        csv_file = self.export_to_csv(results, f"{base_filename}.csv")
        exported_files['csv'] = csv_file
        
        # Excel export
        excel_file = self.export_to_excel(results, f"{base_filename}.xlsx")
        exported_files['excel'] = excel_file
        
        # Word export (if available)
        if DOCX_AVAILABLE:
            try:
                word_file = self.export_to_word(results, f"{base_filename}.docx")
                exported_files['word'] = word_file
            except Exception as e:
                logger.warning(f"Word export failed: {e}")
        
        # PDF export (if available)
        if PDF_EXPORT_AVAILABLE:
            try:
                pdf_file = self.export_to_pdf(results, f"{base_filename}.pdf")
                exported_files['pdf'] = pdf_file
            except Exception as e:
                logger.warning(f"PDF export failed: {e}")
        
        return exported_files

